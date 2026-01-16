from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
import re
from docx import Document
from docxtpl import DocxTemplate
import subprocess
import os
from google.cloud import storage
import tempfile

class ATS:
    def __init__(self):
        self.website = "https://www.weekday.works/resume-screener"

    def get_resume_score(self, jd, file_path):
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-dev-shm-usage"]
            )

            context = browser.new_context()
            page = context.new_page()

            page.goto(self.website, wait_until="domcontentloaded")

            # ---- JD injection ----
            jd_locator = "/html/body/div/main/main/div/div[2]/div[2]/div[1]/div/div/textarea[1]"
            page.locator(f"xpath={jd_locator}").fill(jd)

            # ---- File upload ----
            file_input_locator = "/html/body/div/main/main/div/div[2]/div[2]/div[2]/div[3]/div/input"
            page.set_input_files(
                f"xpath={file_input_locator}",
                file_path
            )

            # ---- Click Screen ----
            screen_btn_locator = '//*[@id="__next"]/main/main/div/div[2]/div[2]/div[5]/div/button'
            page.locator(f"xpath={screen_btn_locator}").click()

            # ---- SCORE ----
            score_locator = "/html/body/div/main/main/div/div[2]/div[3]/div/div[2]/div[2]/div[2]/div/div/div[8]"

            try:
                page.locator(f"xpath={score_locator}").wait_for(timeout=60000)
            except PWTimeout:
                raise Exception("Score not detected")

            score_text = page.locator(f"xpath={score_locator}").inner_text()
            match = re.search(r"\d+", score_text)
            if match:
                score = int(match.group())
            else:
                score = 0

            # ---- REMARKS ----
            remarks_locator = '/html/body/div/main/main/div/div[2]/div[3]/div/div[2]/div[2]/div[2]/div/div/div[9]'
            remarks = page.locator(f"xpath={remarks_locator}").inner_text().strip()

            browser.close()
            return score, remarks

class DocxFunctions():
    def __init__(self):
        pass

    def extract_text_from_docx(self, file_path):
        doc = Document(file_path)
        content = []
        
        # Extract paragraphs with styles
        for para in doc.paragraphs:
            if para.text.strip():
                content.append({
                    'type': 'paragraph',
                    'text': para.text,
                    'style': para.style.name
                })
        
        # Extract tables
        for table in doc.tables:
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            content.append({'type': 'table', 'data': table_data})
        
        return content
    
    def parse_txt_to_string(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text

    def generate_modified_docx(self, template_path, context, output_path):
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)

    def convert_docx_to_pdf(self, docx_path, pdf_path):
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path], check=True)

class FileManager:
    def __init__(self):
        self.client = storage.Client()
        self.bucket = self.client.bucket("job-ref-sys-bucket")
        self.initial_docx_key = "assets/bhavik-h-patel-resume.docx"
        self.initial_pdf_key = "assets/bhavik-h-patel-resume.pdf"

    def get_initial_docx_path(self):
        try:
            blob = self.bucket.blob(self.initial_docx_key)
            temp_file = tempfile.NamedTemporaryFile(
            delete=False,   # keep file after closing
            suffix=".docx",
            mode="wb"
            )
            blob.download_to_file(temp_file)
            output = temp_file.name
            temp_file.close()
        except Exception as e:
            return {
                "error": str(e),
                "message": "Failed to download initial docx",
                "status": "failure"
            }
        return {
            "data": output,
            "status": "success"
        }
    
    def get_initial_pdf_path(self):
        try:
            blob = self.bucket.blob(self.initial_pdf_key)
            temp_file = tempfile.NamedTemporaryFile(
            delete=False,   # keep file after closing
            suffix=".pdf",
            mode="wb"                
            )
            blob.download_to_file(temp_file)
            output = temp_file.name
            temp_file.close()
        except Exception as e:
            return {
                "error": str(e),
                "status": "failure",
                "message": "Failed to download initial pdf"

            }
        return {
            "data": output,
            "status": "success",
        }
        
    def get_template_path(self):
        try:
            blob = self.bucket.blob("assets/resume_template.docx")
            temp_file = tempfile.NamedTemporaryFile(
            delete=False,   # keep file after closing
            suffix=".docx",
            mode="wb"
            )
            blob.download_to_file(temp_file)
            output = temp_file.name
            temp_file.close()
        except Exception as e:
            return {
                "error": str(e),
                "status": "failure",
                "message": "Failed to download template docx"
            }
        return {
            "data": output,
            "status": "success",
        }

    def upload_resume(self, docx_path, pdf_path, job_id):
        try:
            docx_blob = self.bucket.blob(f"resumes/{job_id}/bhavik-h-patel-resume.docx")
            docx_blob.upload_from_filename(docx_path)

            pdf_blob = self.bucket.blob(f"resumes/{job_id}/bhavik-h-patel-resume.pdf")
            pdf_blob.upload_from_filename(pdf_path)

        except Exception as e:
            return {
                "error": str(e),
                "status": "failure",
                "message": "Failed to upload resumes"
            }
        return {
            "status": "success",
        }
    
